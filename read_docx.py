# encoding=utf-8
import os
from docx import Document
from win32com.client import Dispatch
import re
import json

path1 = "E:\\党支部\\关于发放思想政治学习教育及基层组织工作专业知识考学题库的通知(1)\\观澜街道考学题库\\"
path2 = path1+"党的基本理论知识\\"
path3 = path1+"党建业务知识\\"


# 读取目录下所有文件
def findDocxFileList(path):
    list = []
    for root, dirs, files in os.walk(path):
        for file in files:
            # 现有文件名
            localFile = os.path.join(root, file)
            if localFile[-5:] == '.docx':
                print(localFile)
                list.append(localFile)

    print('--------------------', len(list), '------------------')
    return list


def reduceDocx(inpath, outfilepath):
    # 创建写入的word文件对象
    file = Document()
    # 读取path路径下所有.docx文件
    for local in findDocxFileList(inpath):
        # file.add_paragraph('--tce'+local)
        # 读取文件的每一行
        # list = []
        print(local)
        for para in Document(local).paragraphs:
            # 将所有.docx文件内容写入到创建的写入word文件
            file.add_paragraph(para.text)

    # 保存word文件对象
    file.save(outfilepath)

# def filterDocx():
#     path = 'E:\\bigdata\\test.docx'
#     file = Document(path)
#     for para in file.paragraphs:
#         pattern = re.compile(r'_(.*)_',re.I)
#         matchObj = pattern.findall(para.text)
#         if len(matchObj) > 0 :
#             for str in matchObj:
#                 subObj = re.sub('_','',str)
#                 print('前：',str,'----后：',subObj)
    # para.text = subObj


# def mapperDocx(inpath, outfilepath):
#     # 将题目类型分类，并写入对应的word文件中
#     #
#     print(1)

# def readDocx(path):
#     file = Document(path)
#     for para in file.paragraphs:
#         print(para.text)

# 将目录下所有doc文件转成 tpye 文件
# def alldocx2tpye(path,tpye):
#     # 打开word应用程序
#     word = Dispatch('Word.Application')
#     # 后台运行,不显示
#     word.Visible = 0
#     # 不显示提示弹窗
#     word.DisplayAlerts = 0

#     # 获取本地文件路径
#     for root,dirs,files in os.walk(path):
#         # for dir in dirs:
#             # print(os.path.join(root,dir))
#         # print('--------------------------------------')
#         for p in files:
#             print('--------------------------------------')
#             print(p)
#             # 重命名
#             newpath = os.path.splitext(p)[0]+"."+tpye
#             # 用word程序打开word文件
#             doc = word.Documents.Open(p)
#             # 另存为
#             doc.SaveAs(newpath,12,False,"",True,"",False,False,False,False)
#             # 关闭
#             doc.Close()
#             # 删除旧文件
#             # os.remove(p)

#     # 退出word程序
#     word.Quit()

# 读取目录下所有文件
def findDocxFileList(path):
    list = []
    for root, dirs, files in os.walk(path):
        for file in files:
            print('------------', root[root.rfind('\\')+1:], '------------')
            print(file)
            list.append(file)
            # 现有文件名
            # localFile = os.path.join(root, file)
            # if localFile[-5:] == '.docx':
            #     print(localFile)
            #     list.append(localFile)

    print('--------------------', len(list), '------------------')
    return list

# 统一文档格式
# () 替换成中文 （）
# 题目标题等追加回车符：例如 1、选择题 ,将选择题后增加 回车符
# A、x b、x c、x  选择则题答案， 替换成 A.x B.x C.x
# ^l 软回车 替换成 ^p 硬回车
# ' ' 空格替换成 ''
# '．' 替换成 '.'
# 多选题目选项大于 4个的：E、F、G、....
# 类似 __A__
# 答案不在 （）括号内
# A 在最后
# A。 在最后

yun_path = 'cloud://sz-party-building-6qcj9.737a-sz-party-building-6qcj9/doc/'
rowIndex = 1
columnIndex = 1

exam_id = ''
exam_name = ''
book_count = 0

books_json = []
book_json = {}
hasBook = False

q_index = 0
questions_json = []
question_json = {}
q_list = []

q_type = ''
# q_types = ['单选题','填空题','多选题','选择题','判断题','简答题','论述题']
q_types = [
    '单项选择题',
    '多项选择题',
    '多选选择题',
    '判断题',
    '单选题',
    '多选题',
    '选择题',
    '填空题',
    '简答题',
    '论述题'
    ]

def write(xlSheet, value):
   if len(value) > 0:
        if question_json.get('q_id',None) == None:
            global q_index
            q_index += 1
            question_json['q_id'] = 'q_20190610_' +str(q_index)
            question_json['b_id'] = book_json['b_id']
            question_json['content'] = value
            question_json['exam_id'] = [
                exam_id
            ]
        
        isNotType = True
        for tp in q_types:
            if value.find(tp) > -1 :
                global q_type
                q_type = tp
                isNotType = False
                return

        if isNotType:
            if question_json.get('type',None) == None:
                question_json['type'] = q_type
            else:
                q_list.append(value)

        # global columnIndex
        # xlSheet.Cells(rowIndex, columnIndex).Value = value
        # columnIndex += 1



if __name__ == "__main__":
    out = "e:\\bigdata\out.docx"
    local = 'e:\\bigdata\\output.xlsx'

    xl = Dispatch('Excel.Application')

    xl.Workbooks.Open(local)
    xl.Visible = True
    xlBook = xl.Workbooks(1)
    xlSheet = xl.Sheets(1)

    rowIndex = 1
    columnIndex = 1

    for para in Document(out).paragraphs:
        value = para.text
        if len(value) > 0:
            if value == '党建业务知识':
                exam_id = 'exam_2019060901'
                exam_name = '党建业务知识'
                continue
            elif value == '党的基本理论知识':
                exam_id = 'exam_2019060902'
                exam_name = '党的基本理论知识'
                continue
            elif value == 'b_id':
                hasBook = True
                book_count += 1
                book_json = {
                    'b_id':  'b_000' + str(book_count),
                    'name': None,
                    'type': None,
                    'exam_id': [
                        exam_id
                    ]
                }
                continue
            elif book_json['name'] == None:
                book_json['name'] = value
                continue
            elif book_json['type'] == None:
                book_json['type'] = value
                book_json['path'] = yun_path + exam_name+'/' + book_json['name'] +'（'+ value +'）.docx'
                continue

            indexA = value.find('A.')
            indexB = value.find('B.')
            indexC = value.find('C.')
            indexD = value.find('D.')
            indexE = value.find('E.')
            indexF = value.find('F.')
            indexG = value.find('G.')

            if indexA >= 0:
                # 有A
                # 判断是否有B
                if indexB >= 0:
                    # 有B
                    write(xlSheet, value[indexA:indexB])
                    if indexC >= 0:
                        # 有C
                        write(xlSheet, value[indexB:indexC])
                        if indexD >= 0:
                            # 有D
                            write(xlSheet,  value[indexC:indexD])
                            if indexE >= 0:
                                # 有E
                                write(xlSheet,  value[indexD:indexE])
                                if indexF >= 0:
                                    # 有F
                                    write(
                                        xlSheet,  value[indexE:indexF])
                                    if indexG >= 0:
                                        # 有G
                                        write(
                                            xlSheet, value[indexF:indexG])
                                        write(
                                            xlSheet, value[indexG:len(value)])
                                    else:
                                        # 无G
                                        write(
                                            xlSheet, value[indexF:len(value)])
                                else:
                                    # 无F
                                    write(
                                        xlSheet, value[indexE:len(value)])
                            else:
                                # 无E
                                write(
                                    xlSheet, value[indexD:len(value)])
                        else:
                            # 无D
                            write(xlSheet, value[indexC:len(value)])
                    else:
                        # 无C：
                        write(xlSheet, value[indexB:len(value)])
                else:
                    # 无B
                    write(xlSheet, value)
            else:
                # 无A
                # 判断是否有B
                if indexB >= 0:
                    # 有B
                    write(xlSheet, value[0:indexB])
                    if indexC >= 0:
                        # 有C
                        write(xlSheet, value[indexB:indexC])
                        if indexD >= 0:
                            # 有D
                            write(xlSheet, value[indexC:indexD])
                            if indexE >= 0:
                                # 有E
                                write(xlSheet, value[indexD:indexE])
                                if indexF >= 0:
                                    # 有F
                                    write(
                                        xlSheet, value[indexE:indexF])
                                    if indexG >= 0:
                                        # 有G
                                        write(
                                            xlSheet, value[indexF:indexG])
                                        write(
                                            xlSheet, value[indexG:len(value)])
                                    else:
                                        # 无G
                                        write(
                                            xlSheet, value[indexF:len(value)])
                                else:
                                    # 无F
                                    write(
                                        xlSheet, value[indexE:len(value)])
                            else:
                                # 无E
                                write(
                                    xlSheet, value[indexD:len(value)])
                        else:
                            # 无D
                            write(xlSheet, value[indexC:len(value)])
                    else:
                        # 无C：
                        write(xlSheet, value[indexB:len(value)])
                else:
                    # 没B
                    if len(value) > 0:
                        # 保存题目
                        write(xlSheet, value)
                        # 解析答案
                        # if columnIndex == 2:
                        #     # 只有第二排才显示答案，其他不显示
                        #     indexBracketStart = value.rfind('（')
                        #     indexBracketEnd = value.rfind('）')
                        #     if indexBracketStart >= 0 and indexBracketEnd >= 0 and indexBracketEnd > indexBracketStart:
                        #         answer = value[indexBracketStart +
                        #                        1:indexBracketEnd]
                        #         print('index:', indexBracketStart,
                        #               indexBracketEnd)
                        #         print('index value:', answer)
                        #         write(xlSheet, answer)
            
        else:
            # 遇到换行
            if hasBook:
                hasBook = False
                books_json.append(book_json)

            if question_json.get('q_id',None)!=None:
                question_json['list'] = q_list
                questions_json.append(question_json)
                # 复原
                question_json={}
                q_list = []

            rowIndex += 1
            columnIndex = 1

    print('------', len(books_json), '----------')

    file = open('exam_book.json', mode='w+', encoding='utf-8')
    for book in books_json:
        file.write(str(json.dumps(book, ensure_ascii=False)))
        file.write('\n')
    file.close()

    print('------', len(questions_json), '----------')

    file = open('question.json', mode='w+', encoding='utf-8')
    # print(questions_json)
    for ques in questions_json:
        json_str = json.dumps(ques, ensure_ascii=False)
        print(json_str)
        file.write(json_str)
        file.write('\n')

    file.close()
    
    # 遇到回车
    # if count > 0:
    #     # print(count)
    #     # 如果有数据则写入到 excel中
    #     for value in list:
    # indexA = value.find('A.')
    # indexB = value.find('B.')
    # indexC = value.find('C.')
    # indexD = value.find('D.')
    # indexE = value.find('E.')
    # indexF = value.find('F.')
    # indexG = value.find('G.')
    # if indexA >= 0:
    #     # 有A
    #     # 判断是否有B
    #     if indexB >= 0:
    #         # 有B
    #         write(xlSheet, value[indexA:indexB])
    #         if indexC >= 0:
    #             # 有C
    #             write(xlSheet, value[indexB:indexC])
    #             if indexD >= 0:
    #                 # 有D
    #                 write(xlSheet,  value[indexC:indexD])
    #                 if indexE >= 0:
    #                     # 有E
    #                     write(xlSheet,  value[indexD:indexE])
    #                     if indexF >= 0:
    #                         # 有F
    #                         write(
    #                             xlSheet,  value[indexE:indexF])
    #                         if indexG >= 0:
    #                             # 有G
    #                             write(
    #                                 xlSheet, value[indexF:indexG])
    #                             write(
    #                                 xlSheet, value[indexG:len(value)])
    #                         else:
    #                             # 无G
    #                             write(
    #                                 xlSheet, value[indexF:len(value)])
    #                     else:
    #                         # 无F
    #                         write(
    #                             xlSheet, value[indexE:len(value)])
    #                 else:
    #                     # 无E
    #                     write(
    #                         xlSheet, value[indexD:len(value)])
    #             else:
    #                 # 无D
    #                 write(xlSheet, value[indexC:len(value)])
    #         else:
    #             # 无C：
    #             write(xlSheet, value[indexB:len(value)])
    #     else:
    #         # 无B
    #         write(xlSheet, value)
    # else:
    #     # 无A
    #     # 判断是否有B
    #     if indexB >= 0:
    #         # 有B
    #         write(xlSheet, value[0:indexB])
    #         if indexC >= 0:
    #             # 有C
    #             write(xlSheet, value[indexB:indexC])
    #             if indexD >= 0:
    #                 # 有D
    #                 write(xlSheet, value[indexC:indexD])
    #                 if indexE >= 0:
    #                     # 有E
    #                     write(xlSheet, value[indexD:indexE])
    #                     if indexF >= 0:
    #                         # 有F
    #                         write(
    #                             xlSheet, value[indexE:indexF])
    #                         if indexG >= 0:
    #                             # 有G
    #                             write(
    #                                 xlSheet, value[indexF:indexG])
    #                             write(
    #                                 xlSheet, value[indexG:len(value)])
    #                         else:
    #                             # 无G
    #                             write(
    #                                 xlSheet, value[indexF:len(value)])
    #                     else:
    #                         # 无F
    #                         write(
    #                             xlSheet, value[indexE:len(value)])
    #                 else:
    #                     # 无E
    #                     write(
    #                         xlSheet, value[indexD:len(value)])
    #             else:
    #                 # 无D
    #                 write(xlSheet, value[indexC:len(value)])
    #         else:
    #             # 无C：
    #             write(xlSheet, value[indexB:len(value)])
    #     else:
    #         # 没B
    #         if len(value) > 0:
    #             # 保存题目
    #             write(xlSheet, value)
    #             # 解析答案
    #             # if columnIndex == 2:
    #             #     # 只有第二排才显示答案，其他不显示
    #             #     indexBracketStart = value.rfind('（')
    #             #     indexBracketEnd = value.rfind('）')
    #             #     if indexBracketStart >= 0 and indexBracketEnd >= 0 and indexBracketEnd > indexBracketStart:
    #             #         answer = value[indexBracketStart +
    #             #                        1:indexBracketEnd]
    #             #         print('index:', indexBracketStart,
    #             #               indexBracketEnd)
    #             #         print('index value:', answer)
    #             #         write(xlSheet, answer)
    # rowIndex += 1
    # columnIndex = 1
    # count = 0
    # list = []

    #  # 打开excel应用程序


# def openXlsx():
#     local = 'e:\\bigdata\\empty_book.xlsx'
#     xl = Dispatch('Excel.Application')
#     # 后台运行,不显示
#     excel.Visible = 0
#     # # 不显示提示弹窗
#     excel.DisplayAlerts = 0
#     xl.Workbooks.Open(local)
#     xlBook = xl.Workbooks(1)
#     xlSheet = xl.Sheets(1)
    # xlSheet.Cells(1,1).Value = 'What shall be the number of thy counting?'
    # xlSheet.Cells(2,1).Value = 3
    # print(xlSheet.Cells(1,1).Value)
    # print(xlSheet.Cells(2,1).Value)

    # reduceDocx(path1,out)
    # findDocxFileList(path1)
    # error1  = path3+'01中国共产党基层组织选举工作暂行条例题库（单选10）.docx'
    # outfilepath  = 'e:\\bigdata\\newout.docx'
    # a:追加模式
    # 方法一 用docx 读取word输出到txt文件 ，问题：系统编码会丢失~~~
    # outFile = open(out,'a',encoding='utf-8')
    # file = Document(error1)
    # for para in file.paragraphs:
    # outFile.write(para.text+'\n')

    # 方法二 用 win32com 另存为 .txt文件
    # 打开word应用程序
    # word = Dispatch('Word.Application')
    # # 后台运行,不显示
    # word.Visible = 0
    # # 不显示提示弹窗
    # word.DisplayAlerts = 0
    # doc = word.Documents.Open(error1)
    # # newpath = os.path.splitext(error1)[0]+".docx"
    # doc.SaveAs(out,2)
    # doc.Close()
    # word.Quit()

    # for root,dirs,files in os.walk(path):
    #     for p in files:
    #         print('--------------------------------------')
    #         print(p)
    #         newpath = os.path.splitext(p)[0]+".txt"
    #         print(newpath)
